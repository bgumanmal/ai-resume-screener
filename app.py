"""
Two-agent demo: Requirements Extractor -> Resume Screener

Pipeline (handoff pattern, no loop):
    1. Agent 1 (Requirements Extractor) reads the job posting and produces
       a structured checklist of requirements.
    2. Agent 2 (Resume Screener) takes that checklist + the resume and
       checks each item off, then gives an overall recommendation.

This file is the local "server" — a small program that stays running on
your machine and listens for requests from index.html in the same folder.

Setup (run these once, in a terminal, from this folder):
    pip install flask flask-cors python-docx pypdf anthropic python-dotenv

Accepted file types for job posting / resume uploads: .docx and .pdf
(text-based PDFs only — scanned/photographed PDFs have no extractable
text and are not supported; paste the text instead for those).

Credentials live in a .env file in this same folder (see .env.example).
Copy it to .env and fill in your real token — app.py loads it automatically
on startup, so you never have to type it into a terminal.

Then run:
    python app.py

Then open index.html in your browser (double-click it, or right-click ->
Open with -> your browser). Leave this terminal window open while you use
the page — closing it stops the server.
"""

import os
import json
import time
import traceback

from flask import Flask, request, jsonify, Response, stream_with_context
from flask_cors import CORS
from docx import Document
from pypdf import PdfReader
from anthropic import Anthropic
from dotenv import load_dotenv

load_dotenv()  # reads .env in this folder into the environment, if present

app = Flask(__name__)
CORS(app)  # allows index.html (opened as a local file) to call this server

# --- Gateway-aware Anthropic client -----------------------------------
# Reads ANTHROPIC_BASE_URL / ANTHROPIC_AUTH_TOKEN from the environment.
# If those aren't set, falls back to ANTHROPIC_API_KEY talking to
# Anthropic directly (useful if you ever test outside the gateway).
BASE_URL = os.environ.get("ANTHROPIC_BASE_URL")
AUTH_TOKEN = os.environ.get("ANTHROPIC_AUTH_TOKEN")
API_KEY = os.environ.get("ANTHROPIC_API_KEY")

if BASE_URL and AUTH_TOKEN:
    client = Anthropic(base_url=BASE_URL, auth_token=AUTH_TOKEN)
elif API_KEY:
    client = Anthropic(api_key=API_KEY)
else:
    client = None  # will raise a clear error on first request instead of at import time

MODEL = os.environ.get("ANTHROPIC_MODEL", "us.anthropic.claude-sonnet-4-6")

# --- Agent 1: Requirements Extractor ------------------------------------
REQUIREMENTS_SYSTEM_PROMPT = """You are a Requirements Extractor. You are \
given only a job posting (you never see any resume). Your job is to turn \
it into a structured, checkable list of what "fit" means for this role.

Output ONLY a single JSON object, no prose, no markdown fences, with
exactly these fields:
{
  "role_title": "string",
  "requirements": [
    {
      "item": "short, specific, checkable requirement (e.g. '5+ years \
Python backend development')",
      "type": "must_have" | "nice_to_have",
      "category": "skill" | "experience" | "education" | "certification" \
| "other"
    }
  ],
  "hard_constraints": [ "e.g. 'must be based in the US', 'requires active \
security clearance'" ]
}

Extract only what is explicitly stated or clearly implied in the posting.
Do not invent requirements that aren't there."""

# --- Agent 2: Resume Screener --------------------------------------------
SCREENER_SYSTEM_PROMPT = """You are a Resume Screener. You are given a \
structured requirements checklist (produced by another agent from the job \
posting) and a candidate's resume. Go through the checklist item by item.

Output ONLY a single JSON object, no prose, no markdown fences, with
exactly these fields:
{
  "role_title": "copied from the checklist",
  "item_results": [
    {
      "item": "copied from the checklist",
      "type": "must_have" | "nice_to_have",
      "status": "met" | "partially_met" | "not_met" | "unclear_from_resume",
      "evidence": "one line quoting or pointing to what in the resume \
justifies this status, or 'not mentioned in resume' if absent"
    }
  ],
  "hard_constraints_check": [
    { "constraint": "copied from checklist", "status": "met" | "not_met" \
| "unclear_from_resume", "evidence": "one line" }
  ],
  "overall_recommendation": "strong_fit" | "possible_fit" | "not_a_fit",
  "recommendation_reason": "2-3 sentences justifying the overall call, \
referencing which must-haves drove it"
}

Only use evidence actually present in the resume text. Do not assume
skills or experience that aren't stated. Be precise, not generous."""

# --- Agent 3: Fitment Reviser (human-in-the-loop recompute) --------------
# Not part of the original pipeline — this runs only when an HR reviewer
# overrides one or more item statuses via the dropdowns in the UI and
# clicks "Revise Fitment Analysis". It does NOT re-examine the resume; it
# only re-weighs the (possibly human-corrected) item statuses it's given
# and rewrites the overall verdict + reasoning to match. This keeps the
# human's judgment as the source of truth for anything they changed, while
# still producing a natural-language rationale instead of a mechanical
# rule ("no explanation why").
REVISER_SYSTEM_PROMPT = """You are a Fitment Reviser. An HR reviewer has \
looked at an AI-generated checklist screening and, using their own \
judgment, may have changed the status of one or more items (for example, \
turning a "not_met" into "met" because they know about a certification \
that wasn't in the AI's evidence, or the reverse).

You are given the full list of checklist items with their CURRENT status \
(after any human edits) and a flag on which ones were human-edited, plus \
the hard constraints checklist. Do not second-guess or revert the human's \
edits — treat every current status, human-edited or not, as ground truth. \
Your only job is to recompute the overall verdict and write a fresh \
rationale that reflects the CURRENT statuses, explicitly noting where a \
human correction changed the picture.

Output ONLY a single JSON object, no prose, no markdown fences, with
exactly these fields:
{
  "overall_recommendation": "strong_fit" | "possible_fit" | "not_a_fit",
  "recommendation_reason": "2-3 sentences justifying the revised call, \
referencing which must-haves drove it and explicitly mentioning any \
human-edited item that changed the outcome"
}"""


def extract_text_from_docx(file_storage) -> str:
    """Pull plain text out of an uploaded .docx file."""
    doc = Document(file_storage)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_pdf(file_storage) -> str:
    """Pull plain text out of an uploaded text-based PDF.

    This only works for PDFs where the text is stored as real characters
    (e.g. exported from Word, LinkedIn, a printer driver, etc.). A scanned
    or photographed PDF has no extractable text — pypdf will just return
    little or nothing for those, which we detect and flag below rather
    than silently sending an empty document to the agent."""
    reader = PdfReader(file_storage)
    parts = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(parts).strip()

    if len(text) < 20:
        raise ValueError(
            "this PDF has little or no extractable text — it's likely a "
            "scanned/image-based PDF, which isn't supported. Paste the "
            "text instead, or use a .docx."
        )
    return text


def get_text_input(field_name: str) -> str:
    """For a given field ('job' or 'resume'), prefer pasted text if present,
    otherwise extract text from an uploaded .docx or .pdf file."""
    pasted = request.form.get(f"{field_name}_text", "").strip()
    if pasted:
        return pasted

    file_key = f"{field_name}_file"
    if file_key in request.files and request.files[file_key].filename:
        f = request.files[file_key]
        name = f.filename.lower()
        if name.endswith(".docx"):
            return extract_text_from_docx(f)
        elif name.endswith(".pdf"):
            try:
                return extract_text_from_pdf(f)
            except ValueError as e:
                raise ValueError(f"{field_name}: {e}")
        else:
            raise ValueError(f"{field_name}: only .docx and .pdf files are supported, got '{f.filename}'")

    raise ValueError(f"{field_name}: no text pasted and no file uploaded")


def call_agent(system_prompt: str, user_message: str, max_tokens: int = 1800) -> dict:
    if client is None:
        raise RuntimeError(
            "No Anthropic credentials found. Set ANTHROPIC_BASE_URL + "
            "ANTHROPIC_AUTH_TOKEN (gateway) or ANTHROPIC_API_KEY (direct) "
            "as environment variables before starting this server."
        )
    response = client.messages.create(
        model=MODEL,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )
    raw = response.content[0].text.strip()
    # Models sometimes wrap JSON in ```json fences despite instructions —
    # strip those defensively so json.loads doesn't choke on a real run.
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.lower().startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            f"the agent's reply was cut off because it hit the {max_tokens}-token "
            f"limit before finishing — the checklist or resume is likely longer than "
            f"this demo's default budget. Increase max_tokens in call_agent() for this "
            f"call and try again."
        )

    return json.loads(raw)


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "model": MODEL, "gateway_configured": client is not None})


@app.route("/process", methods=["POST"])
def process():
    try:
        job_text = get_text_input("job")
        resume_text = get_text_input("resume")
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    try:
        # --- Agent 1 ---
        requirements = call_agent(
            REQUIREMENTS_SYSTEM_PROMPT,
            f"Job posting:\n{job_text}",
        )

        # --- Agent 2 ---
        # Uses a higher max_tokens than Agent 1: it has to write one
        # evidence-backed row per checklist item, which can add up for
        # postings with many requirements.
        screening = call_agent(
            SCREENER_SYSTEM_PROMPT,
            f"Requirements checklist:\n{json.dumps(requirements, indent=2)}\n\n"
            f"Candidate resume:\n{resume_text}",
            max_tokens=3000,
        )

        return jsonify({
            "requirements": requirements,
            "screening": screening,
        })

    except json.JSONDecodeError as e:
        return jsonify({"error": f"Agent returned non-JSON output, could not parse: {e}"}), 502
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"{type(e).__name__}: {e}"}), 500


@app.route("/process_stream", methods=["POST"])
def process_stream():
    """Same pipeline as /process, but reports progress as it goes.

    Streams one JSON object per line (newline-delimited JSON — not quite
    SSE, since we're driven by a POST + file upload rather than a plain
    GET EventSource). Each line looks like:
        {"stage": "agent1", "status": "start"}
        {"stage": "agent1", "status": "done", "elapsed": 3.21, "data": {...}}
    so the page can light up each box of the pipeline diagram, show a
    running timer while a stage is "start" but not yet "done", and record
    the elapsed seconds once it finishes.

    Reading the uploaded files / form fields has to happen BEFORE the
    generator starts (Flask's `request` isn't available once the response
    has started streaming), so we grab job_text/resume_text up front and
    only stream the parts that take real time: preprocessing already
    happened by the time we get here, so we time it separately using a
    wall-clock timestamp taken right at the top of this view function.
    """
    stream_start = time.time()
    preprocess_error = None
    job_text = resume_text = None
    try:
        job_text = get_text_input("job")
        resume_text = get_text_input("resume")
    except ValueError as e:
        preprocess_error = str(e)
    preprocess_elapsed = round(time.time() - stream_start, 2)

    def emit(obj):
        return json.dumps(obj) + "\n"

    def generate():
        # --- Stage 1: pre-processing (file extraction / text prep) ---
        yield emit({"stage": "preprocess", "status": "start"})
        if preprocess_error:
            yield emit({"stage": "preprocess", "status": "error", "message": preprocess_error, "elapsed": preprocess_elapsed})
            return
        yield emit({"stage": "preprocess", "status": "done", "elapsed": preprocess_elapsed})

        # --- Stage 2: Agent 1 — Requirements Extractor ---
        yield emit({"stage": "agent1", "status": "start"})
        t0 = time.time()
        try:
            requirements = call_agent(
                REQUIREMENTS_SYSTEM_PROMPT,
                f"Job posting:\n{job_text}",
            )
        except json.JSONDecodeError as e:
            yield emit({"stage": "agent1", "status": "error", "message": f"Agent returned non-JSON output: {e}", "elapsed": round(time.time() - t0, 2)})
            return
        except Exception as e:
            traceback.print_exc()
            yield emit({"stage": "agent1", "status": "error", "message": f"{type(e).__name__}: {e}", "elapsed": round(time.time() - t0, 2)})
            return
        agent1_elapsed = round(time.time() - t0, 2)
        yield emit({"stage": "agent1", "status": "done", "elapsed": agent1_elapsed, "data": requirements})

        # --- Stage 3: Agent 2 — Resume Screener ---
        yield emit({"stage": "agent2", "status": "start"})
        t1 = time.time()
        try:
            screening = call_agent(
                SCREENER_SYSTEM_PROMPT,
                f"Requirements checklist:\n{json.dumps(requirements, indent=2)}\n\n"
                f"Candidate resume:\n{resume_text}",
                max_tokens=3000,
            )
        except json.JSONDecodeError as e:
            yield emit({"stage": "agent2", "status": "error", "message": f"Agent returned non-JSON output: {e}", "elapsed": round(time.time() - t1, 2)})
            return
        except Exception as e:
            traceback.print_exc()
            yield emit({"stage": "agent2", "status": "error", "message": f"{type(e).__name__}: {e}", "elapsed": round(time.time() - t1, 2)})
            return
        agent2_elapsed = round(time.time() - t1, 2)
        yield emit({"stage": "agent2", "status": "done", "elapsed": agent2_elapsed, "data": screening})

        # --- Stage 4: Decision handoff — final assembly, no extra model call ---
        yield emit({"stage": "decision", "status": "start"})
        total_elapsed = round(time.time() - stream_start, 2)
        yield emit({
            "stage": "decision",
            "status": "done",
            "elapsed": 0.0,
            "total_elapsed": total_elapsed,
            "data": {"requirements": requirements, "screening": screening},
        })

    return Response(stream_with_context(generate()), mimetype="application/x-ndjson")


@app.route("/revise", methods=["POST"])
def revise():
    """Human-in-the-loop recompute: HR has edited one or more item statuses
    in the browser (dropdowns), and clicked "Revise Fitment Analysis". We
    don't re-read the resume or the job posting here — only the current
    item statuses (as the human has left them) are sent in, so the human's
    edits are always the source of truth for what "met" vs "not_met" means.
    This endpoint only asks the model to recompute the overall verdict and
    write a rationale that matches.

    Expected JSON body:
    {
      "role_title": "...",
      "item_results": [ {item, type, status, evidence, human_edited: bool}, ... ],
      "hard_constraints_check": [ {constraint, status, evidence}, ... ]
    }
    """
    body = request.get_json(silent=True) or {}
    item_results = body.get("item_results", [])
    hard_constraints_check = body.get("hard_constraints_check", [])

    if not item_results:
        return jsonify({"error": "no item_results provided — nothing to revise"}), 400

    user_message = (
        "Checklist items with current status (human_edited=true means an "
        "HR reviewer manually set this one, overriding the AI's original "
        "call):\n"
        f"{json.dumps(item_results, indent=2)}\n\n"
        "Hard constraints, current status:\n"
        f"{json.dumps(hard_constraints_check, indent=2)}"
    )

    try:
        revised = call_agent(REVISER_SYSTEM_PROMPT, user_message, max_tokens=600)
        return jsonify(revised)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Agent returned non-JSON output, could not parse: {e}"}), 502
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"{type(e).__name__}: {e}"}), 500


if __name__ == "__main__":
    print(f"\nModel: {MODEL}")
    print(f"Gateway configured: {client is not None}")
    if client is None:
        print("WARNING: no credentials set — /process will return an error until you set "
              "ANTHROPIC_BASE_URL + ANTHROPIC_AUTH_TOKEN (or ANTHROPIC_API_KEY).\n")
    print("Starting server on http://localhost:5000 ...")
    print("Open index.html in your browser now. Leave this window open.\n")
    app.run(host="127.0.0.1", port=5000, debug=False)
